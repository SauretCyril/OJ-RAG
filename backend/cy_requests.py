from flask import Blueprint, request, jsonify
import logging
import pythoncom
import os
from docx2pdf import convert
from docx import Document
import openai
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
import requests

'''My script'''
from cy_mistral import get_mistral_answer, mistral  # Import the function and Blueprint
from cy_paths import GetRoot



# Configure logger
logging.basicConfig(level=logging.ERROR, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

cy_requests = Blueprint('requests', __name__)

def extract_text_from_pdf(pdf_path):
    try:
        if os.path.exists(pdf_path):
            ##print(f"-> le {pdf_path} fichier existe")
            text= ""
            
            # Utiliser PyPDF2 pour extraire le texte
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            if not text.strip():
                text +="Le PDF est vide ou ne contient pas de texte extractible"
                return text
            
            # Ensure the text is properly encoded and decoded
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            
            #print(f"-> Text extrait: {text[:200]}...")
            return text
            
        #print(". Le fichier PDF n'existe pas")
    except Exception as e:
        print(f"An error occurred while extracting text from PDF: {e}")
        return ""


@cy_requests.route('/get_AI_answer', methods=['POST'])

async def get_job_answer():
    try:
        file = request.json.get('path')
        RQ = request.json.get('RQ')
        NumDos = request.json.get('NumDos')
        ##
        ## load the request from the file if it exists
        ##  
        print (f"dbg 21547 : RQ : {NumDos}")
        the_request = await load_AI_Instructions(".ask", NumDos)
        
        if the_request != "":
            RQ = the_request
        if not file or not RQ:
            logger.error("Er005.Missing job file path or question")
            return jsonify({'Er005': 'Missing job file path or question'}), 400

        text1 = extract_text_from_pdf(file)
        
        if not text1:
            logger.error("Er006.Job text extraction failed")
            return jsonify({'Er006': 'Job text extraction failed'}), 500

        role="En tant qu' expert en analyse d'offres d'emploi dans le domaine informatique (développeur, Analyste ou Testeur logiciel) , analyse le texte suivant et réponds à cette question"
        ## Load the role from the file if it exists
        the_role = await load_AI_Instructions(".role", NumDos)
        
        if the_role != "":
            role = the_role
        answer = get_mistral_answer(RQ,role,text1)
        return jsonify({
            'raw_text': text1,
            'formatted_text': answer
        })

    except Exception as e:
        logger.error(f"Er007.Error: {str(e)}")
        return jsonify({'Er007': str(e)}), 500

@cy_requests.route('/save-answer', methods=['POST'])
def save_answer():
    try:
        job_text_data = request.json.get('text_data')
        job_number = request.json.get('number')
        the_path = request.json.get('the_path')
        rq = request.json.get('RQ')
        if the_path == '':
            the_path = GetRoot()

        if not job_text_data or not job_number:
            logger.error(f"Er008.error.Missing job text data or job number: job_text_data={job_text_data}, job_number={job_number}")
            return jsonify({'Er008': 'Missing job text data or job number'}), 400

        if not the_path:
            logger.error("Er009.Received path is None")
            return jsonify({'Er009': 'Missing path'}), 400

        file_name = f"{job_number}_gpt_request"
        file_path = os.path.join(the_path, job_number)
        if not os.path.exists(file_path):
            #os.makedirs(file_path)
        
            file_path_docx = os.path.join(file_path, file_name + ".docx")
            file_path_RQ = os.path.join(file_path, file_name + "_RQ.txt")

            doc = format_text_as_word_style(job_text_data, job_number)
            doc.save(file_path_docx)

            pdf_file_path = file_path_docx.replace('.docx', '.pdf')
       
        # Gérer la conversion de façon plus robuste
            pythoncom.CoInitialize()  # Initialize COM library
            try:
                # Ajouter un timeout pour éviter que Word reste bloqué
                import time
                start_time = time.time()
                
                # Convertir le fichier
                convert(file_path_docx, pdf_file_path)
                
                # Si la conversion a réussi et que le fichier PDF existe, supprimer le DOCX
                if os.path.exists(pdf_file_path):
                    os.remove(file_path_docx)
                    
            except Exception as convert_error:
                logger.error(f"Er009a.Error during conversion: {str(convert_error)}")
                # Essayer de tuer toutes les instances de Word qui pourraient être bloquées
                try:
                    import subprocess
                    subprocess.run(['taskkill', '/f', '/im', 'WINWORD.EXE'], shell=True)
                except:
                    pass
                raise convert_error
            finally:
                # S'assurer que COM est bien désinitializé
                pythoncom.CoUninitialize()

            # Sauvegarder la requête dans un fichier texte
            save_rq_to_text_file(file_path_RQ, rq)
            
            return jsonify({'dbg009': 'Job text saved successfully', 'pdf_file_path': pdf_file_path})

    except Exception as e:
        logger.error(f"Er009.Error saving job text: {str(e)}")
        return jsonify({'Er009': str(e)}), 500

def save_rq_to_text_file(file_path, rq):
    # Vérifier si le répertoire parent existe
    parent_dir_path = os.path.dirname(file_path)
    if not os.path.exists(parent_dir_path):
        print(f"Création du répertoire manquant: {parent_dir_path}")
        #os.makedirs(parent_dir_path, exist_ok=True
        with open(file_path, 'w') as file:
            file.write(rq)

def format_text_as_word_style(job_text, job_number):
    doc = Document()
    doc.add_heading(f"Job Number: {job_number}", level=1)
    
    for line in job_text.split('\n'):
        if line.startswith('- '):
            doc.add_paragraph(line, style='ListBullet')
        else:
            doc.add_paragraph(line)
    
    return doc


@cy_requests.route('/get_AI_answer_from_url', methods=['POST'])
def get_job_answer_from_url():
    try:
        url = request.json.get('url')
        RQ = request.json.get('RQ')

        if not url or not RQ:
            logger.error("Er014.Missing job file path or question")
            return jsonify({'Er014': 'Missing job file path or question'}), 400

        text1 = extract_text_from_url(url)
        text1 += "<-"+url+"->"
        
        if not text1:
            logger.error("Er0016.Job text extraction failed")
            return jsonify({'Er016': 'Job text extraction failed'}), 500
        role = "En tant qu' expert en analyse d'offres d'emploi dans le domaine informatique (développeur, Analyste ou Testeur logiciel), analyse le texte suivant et réponds à cette question"
        answer = get_mistral_answer(RQ, role, text1)

        return jsonify({
            'raw_text': text1,
            'formatted_text': answer
        })

    except Exception as e:
        logger.error(f"Er018.Error: {str(e)}")
        return jsonify({'Er018': str(e)}), 500





def extract_text_from_url(url):
    try:
        #print(f"-> Tentative d'extraction de texte depuis l'URL: {url}")
        
        # Récupérer le contenu de l'URL
        response = requests.get(url)
        #print(f"-> Statut de la réponse HTTP: {response.status_code}")
        response.raise_for_status()  # Vérifie si la requête a réussi
        
        # Parser le HTML
        soup = BeautifulSoup(response.text, 'html.parser')
        #print("-> HTML parsé avec succès")
        
        # Supprimer les scripts et styles
        for script in soup(["script", "style"]):
            script.decompose()
        #print("-> Scripts et styles supprimés")
            
        # Extraire le texte
        text = soup.get_text(separator='\n')
        #print("-> Texte extrait du HTML")
        
        # Nettoyer le texte (supprimer les lignes vides multiples)
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)
        #print("-> Texte nettoyé")
        
        if not text.strip():
            print("-> L'URL ne contient pas de texte extractible")
            return "L'URL ne contient pas de texte extractible"
            
        print(f"dbg 6523 -> Texte extrait de l'URL: {text[:200]}...")
        return text
        
    except Exception as e:
        print(f"Erreur lors de l'extraction depuis l'URL: {str(e)}")
        return f"Une erreur s'est produite: {str(e)}"


def extract_text(source, is_url=False):
    """
    Extrait le texte soit d'un PDF soit d'une URL
    """
    if is_url:
        return extract_text_from_url(source)
    else:
        return extract_text_from_pdf(source)


def get_answer(question, role,context=""):
    try:
        
        client = openai  # Ensure the openai library is correctly used
        full_context = f"""{role}: {question}\n\nContexte:\n{context}"""
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": role},
                {"role": "user", "content": full_context}
            ],
            temperature=0.8,
            max_tokens=1100
        )
        
        return response.choices[0].message.content

    except Exception as e:
        print(f"Erreur lors de l'analyse: {str(e)}")
        return f"Une erreur s'est produite: {str(e)}"


def favicon():
    return ""


'''
get info of pdf file and return'''
def get_info(file_path,role, question):
    try:
        
        # client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))  # Assurez-vous que OPENAI_API_KEY est défini dans vos variables d'environnement
        # context = extract_text_from_pdf(file_path)
        # if (context == ""):
        #     return "{'url':'', 'entreprise':'inconnue', 'poste':'Annonce non lisible'}"
        
        # #print(f"Contexte extrait: {context[:200]}...")
        # full_context = f"{question}\n\nContexte:\n{context}"
        
        # response = client.chat.completions.create(
        #     model="gpt-3.5-turbo",
        #     messages=[
        #         {"role": "system", "content": "analyse le texte suivant et réponds à cette question, peux tu renvoyer les informations sous forme de données json, les champs son définie dans la question entre [ et ]"},
        #         {"role": "user", "content": full_context}
        #     ],
        #     temperature=0.7,
        #     max_tokens=1000
        # )
        context = extract_text_from_pdf(file_path)
        response = get_mistral_answer(question, role, context)
        return response

    except Exception as e:
        print(f"Erreur lors de l'analyse: {str(e)}")
        return f"Une erreur s'est produite: {str(e)}"




def response_me(question,url,role):
    try:
        
        client = openai  # Assurez-vous que OPENAI_API_KEY est défini dans vos variables d'environnement
        context=extract_text_from_url(url)
        if (context == ""):
            return "{'url':'', 'entreprise':'inconnue', 'poste':'Annonce non lisible'}"
        
        
        full_context = f"{question}\n\nContexte:\n{context}"
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": role},
                {"role": "user", "content": full_context}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        
        return response.choices[0].message.content

    except Exception as e:
        print(f"Erreur lors de l'analyse: {str(e)}")
        return f"Une erreur s'est produite: {str(e)}"

async def load_AI_Instructions(file_name, NumDos):
    try:
        text = ""
        file_name_txt = file_name
        
        dossier = os.path.join(GetRoot(),  NumDos)
        specif_filepath = os.path.join(dossier,file_name_txt)
        print("dbg4578 :fichier requete",specif_filepath)
        
        if os.path.exists(specif_filepath):
            filepath =specif_filepath 
        else:
            default_filepath = os.path.join(GetRoot(), file_name_txt)  # Updated to call GetRoot() correctly
            default_filepath = default_filepath.replace('\\', '/')
            filepath=default_filepath
        #print("dbg3434 :fichier requete",filepath)
        #print("dbg789 :fichier instructions",filepath)
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as file:
                text = file.read()
        
        return text

    except Exception as e:
        logger.error(f"Error loading text: {str(e)}")
        return ""