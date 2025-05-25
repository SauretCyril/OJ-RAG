from flask import Blueprint, request, jsonify, send_from_directory, render_template_string
import os
import json
import platform
import csv
import subprocess
from werkzeug.utils import secure_filename

# Import de la configuration centralisée
from cy_app_config import app_config

import tkinter as tk
from tkinter import filedialog

from docx2pdf import convert
import pythoncom
from fpdf import FPDF
from docx import Document

import json
from tqdm import tqdm
from PyPDF2 import PdfFileReader, PdfFileWriter
from io import BytesIO
import logging
from logging import DEBUG

from dotenv import load_dotenv

from cy_paths import *
from cy_cookies import *
from cy_requests import extract_text_from_pdf
from cy_mistral import get_mistral_answer



load_dotenv()
cy_routes = Blueprint('cy_routes', __name__)
logging.basicConfig(level=DEBUG, format='%(asctime)s - (name)s - (levelname)s - (message)s')
logger = logging.getLogger(__name__)

# Define excluded directories
EXCLUDED_DIRECTORIES = ["suivi", "pile", "conf"]  # Add your excluded directories here
files_type=[{"suffix":"ANNONCE_SUFFIX","type":"AN"}]

# Load shared constants
def load_constants():
    config_path = os.path.join(os.path.dirname(__file__), 'config', 'constants.json')
    with open(config_path, 'r') as f:
        return json.load(f)

CONSTANTS = load_constants()


#print('Loaded constants:', CONSTANTS)


@cy_routes.route('/get_constants', methods=['GET'])
def get_constants():
    # Utilisation de app_config pour récupérer les constantes
    return jsonify(app_config.constants)


from datetime import datetime

def calculate_delay(data):
    try:
        today = datetime.today()
        date_from = data.get('date_from', '')
        date_rep = data.get('date_rep', '')
        todo = data.get('todo', '')

        if 'refus' in todo.lower():
            return "dead"

        if date_rep:
            try:
                date_rep_c = datetime.strptime(date_rep, '%d-%m-%Y')
                #print("dbg778 => date_from_c = "+date_from_c)
                return (today - date_rep_c).days
            except ValueError:
                print ("err dbg456= " +  ValueError )
                pass
            
        if date_from:
            try:
                date_from_c = datetime.strptime(date_from, '%d-%m-%Y')
                #print("dbg778 => date_from_c = "+date_from_c)
                return (today - date_from_c).days
            except ValueError:
                print ("err dbg457= " +  ValueError )
                pass

       

        return 'N/A'
    except Exception as e:
        print(f"Error calculating delay: {e}")
        return 'N/A'


    


@cy_routes.route('/read_annonces_json', methods=['POST'])
async def read_annonces_json():
    try:
        print("DBG-4658.0: Début de la fonction read_annonces_json")
        isDetectNew="O"
        #buildAllPaths()
        data = request.get_json()
        if not data:
            print("DBG-4658.1: Pas de données JSON reçues")
            return jsonify([]), 200  # Retourne un tableau vide en cas de données manquantes
            
        #excluedFile = data.get('excluded', 'excluded_annonces.json')
        excluedFile =".exclued"
        print(f"DBG-4658.2: Fichier d'exclusion: {excluedFile}")
        
        directory_path = GetRoot()
        print(f"DBG-4658.3: Chemin racine: {directory_path}")
       
        if not os.path.exists(GetRoot()):
            print("DBG-4658.4: Le chemin racine n'existe pas")
            return jsonify([]), 200
            
        dossier_list = []
        crit_annonces = None
        # Charger les critères d'annonces de façon sécurisée
        try:
            crit_annonces = load_crit_annonces(excluedFile)
            print(f"DBG-4658.5: Critères d'annonces chargés: {crit_annonces is not None}")
        except Exception as e:
            print(f"ERR-4658.6: Erreur lors du chargement des critères d'annonces: {e}")
            
       
        print (f"DBG-4658-isDetectNew == {isDetectNew}")    
        record_added=False
        print (f"DBG-4658-starting loading-------------------------------------")    
        # Parcourir tous les dossiers sauf ceux exclus
        for root, _, files in os.walk(directory_path):
            parent_dir = os.path.basename(root)
            if parent_dir in EXCLUDED_DIRECTORIES:
                continue
            
            # Continuer le traitement normal...
            # ...existing code...
            file_doc = parent_dir + CONSTANTS['FILE_NAMES']['ANNONCE_SUFFIX'] + ".pdf"
            file_isGptResum = parent_dir + CONSTANTS['FILE_NAMES']['GPT_REQUEST_SUFFIX']
            file_isGptResum_Path1 = os.path.join(root, file_isGptResum)
            file_isGptResum_Path1 = file_isGptResum_Path1.replace('\\', '/')
            file_cv = parent_dir + CONSTANTS['FILE_NAMES']['CV_SUFFIX'] + ".docx"
            file_cv_pdf = parent_dir + CONSTANTS['FILE_NAMES']['CV_SUFFIX'] + ".pdf"
            file_BA_pdf = parent_dir + CONSTANTS['FILE_NAMES']['BA_SUFFIX_NAME'] + ".pdf"
            file_BA_docx = parent_dir + CONSTANTS['FILE_NAMES']['BA_SUFFIX_NAME'] + ".docx"
            
            file_cv_New = parent_dir + CONSTANTS['FILE_NAMES']['CV_SUFFIX_NEW'] + ".docx"
            file_cv_pdf_New = parent_dir + CONSTANTS['FILE_NAMES']['CV_SUFFIX_NEW'] + ".pdf"
            data_json_file = ".data.json" 
            
            data = {}
            isCVin="N"
            isBAdocx="N"
            isCVinpdf="N"
            isBAinpdf="N"
            isAction="N"
            isJo="N"
          
            file_path_isJo=""
            isGptResum="N"  # Initialiser à N par défaut
            
            isJoTyp=""
            
            # Vérification des fichiers présents
            for filename in files:   
                if filename == file_BA_docx: 
                    isBAdocx="O"
                if filename == file_cv or filename == file_cv_New:
                    isCVin="O"
                if filename == file_cv_pdf or filename == file_cv_pdf_New:
                    isCVinpdf="O"
                if filename == file_BA_pdf:
                    isBAinpdf="O"
                if filename == file_doc:
                    file_path_isJo = os.path.join(root, file_doc.replace('\\', '/'))
                    isJo="O"
                if filename == file_isGptResum:
                    isGptResum="O"
                    file_path_gpt = os.path.join(root, file_isGptResum)
            record_added = False
            # Traitement des fichiers de données
            for filename in files:
                file_path = os.path.join(root, filename)
                file_path = file_path.replace('\\', '/')
                file_path_nodata = os.path.join(root, ".data.json")
                file_path_nodata = file_path_nodata.replace('\\', '/')
                
                if filename == data_json_file:
                    record_added = True
                    try:
                        with open(file_path, 'r', encoding='utf-8') as file:
                            datajson = json.load(file)
                             # Vérification des exclusions
                            isExclued = False
                            if datajson["etat"]=="DELETED" or datajson["etat"]=="ARCHIVE" or datajson["etat"]=="ClOSED" or datajson["etat"]=="DONE":
                                isExclued = True
                            
                            
                            if not isExclued:
                                # Ajout des métadonnées supplémentaires
                                datajson["dossier"] = parent_dir
                                datajson["isJo"] = isJo
                                datajson["GptSum"] = isGptResum
                                datajson["CV"] = isCVin
                                datajson["CVpdf"] = isCVinpdf
                                datajson["BA"] = isBAdocx
                                datajson["isAction"] = isAction
                                datajson["BApdf"] = isBAinpdf
                                # Ajouter à la liste des dossiers
                                jData = {file_path: datajson}
                                dossier_list.append(jData)
                                
                                print (f"{parent_dir}-NEW-4658 LOADING : Fichier {file_path} chargé avec succès")
                         
                            
                    except json.JSONDecodeError as e:
                        print(f"{parent_dir}-ERR-4658.a1: Le fichier {file_path} contient du JSON invalide: {str(e)}")
                    except Exception as e:
                        print(f"{parent_dir}-ERR-4658.a2: Erreur lors du traitement de {file_path}: {str(e)}")
           
                     #ici
            # Si aucun enregistrement n'a été ajouté pour ce dossier
            if isDetectNew == "O" :
                
                if not record_added:
                    try:
                        print ("---------------------------")
                        print (f"{parent_dir}-NEW-4658.v - Nouveau dossier")
                        thefile=""
                        Piece_exist=False
                        data = define_default_data()  # Initialize with default values
                        
                        if isJo =="O":
                            thefile= file_path_isJo
                            print (f"{parent_dir}NEW-4658a file_path_isJo trouvé = ",file_path_isJo)
                            Piece_exist=True
                        elif isGptResum =="O":
                            thefile =file_path_gpt
                            Piece_exist=True
                        
                        if Piece_exist:
                            print (f"{parent_dir}NEW-4658b le fichier main va être traité = ",thefile)
                            #thefile = thefile.replace('\\', '/')
                            texte=extract_text_from_pdf(thefile)
                            infos=texte
                            the_request = await load_Instruction_classement()
                            print(f"{parent_dir}NEW-4658c- la question pour le classement", the_request)
                            if not the_request or the_request.strip() == "":
                                print(f"{parent_dir}ERR-4658d: the_request is invalid or empty.")
                                #return jsonify({"status": "error", "message": "Invalid instruction request"}), 400
                                infos=texte
                            else:
                                role="analyse le texte suivant et réponds à cette question, peux tu renvoyer les informations sous forme de données json, les champs son définie dans la question entre [ et ]"
                                print(f"{parent_dir}ERR-4658d : le rôle pour le classement", role)
                                infos = get_mistral_answer(the_request, role, texte)
                                print(f"{parent_dir}NEW-4658e answer mistral = ", infos)
                            if infos:
                                try:
                                    # Tenter de parser comme JSON
                                    parsed_json = json.loads(infos)
                                    data["url"] = parsed_json.get("url", "N/A")
                                    data["Date"] = parsed_json.get("Date", "N/A") 
                                    data["entreprise"] = parsed_json.get("entreprise", "N/A")
                                    data["description"] = parsed_json.get("poste", "N/A")
                                    data["Lieu"] = parsed_json.get("lieu", "N/A")
                                except json.JSONDecodeError:
                                    # La réponse n'est pas du JSON valide
                                    print(f"{parent_dir}-ERR-4658f : Réponse non JSON, tentative d'extraction des infos du texte")
                                    # Extraction basique (peut être améliorée)
                                    try:
                                        # Tenter de trouver des données structurées dans la réponse texte
                                        import re
                                        
                                        # Chercher un objet JSON dans la réponse
                                        json_match = re.search(r'(\{.*\})', infos, re.DOTALL)
                                        if json_match:
                                            try:
                                                extracted_json = json.loads(json_match.group(1))
                                                data["url"] = extracted_json.get("url", "N/A")
                                                data["Date"] = extracted_json.get("Date", "N/A")
                                                data["entreprise"] = extracted_json.get("entreprise", "N/A")
                                                data["description"] = extracted_json.get("poste", "N/A")
                                                data["Lieu"] = extracted_json.get("lieu", "N/A")
                                            except:
                                                # Utiliser le texte brut
                                                data["url"] = "N/A"
                                                data["Date"] = "N/A" 
                                                data["entreprise"] = "N/A"
                                                data["description"] = "Pas d'infos"
                                                data["Lieu"] = "N/A"
                                        else:
                                            # Utiliser le texte brut
                                            data["url"] = "N/A"
                                            data["Date"] = "N/A" 
                                            data["entreprise"] = "N/A"
                                            data["description"] = "Pas d'infos"
                                            data["Lieu"] = "N/A"
                                    except Exception as extraction_error:
                                        print(f"{parent_dir}-ERR-4658g : Erreur lors de l'extraction: {str(extraction_error)}")
                                        # Fallback en cas d'échec total
                                        data["url"] = "N/A"
                                        data["Date"] = "N/A"
                                        data["entreprise"] = "N/A"  
                                        data["description"] = "Erreur lors du traitement"
                                        data["Lieu"] = "N/A"
                                
                            # Préparation des données
                            data["dossier"] = parent_dir    
                            data["isJo"] = isJo
                            data["isAction"] = isAction
                            data["GptSum"] = isGptResum
                            data["CV"] = isCVin
                            data["CVpdf"] = isCVinpdf
                            data["BA"] = isBAdocx
                            data["BApdf"] = isBAinpdf
                            # block info piece         
                            data["etat"] = "New"
                              # Créer l'objet de données pour ajouter à la liste
                            file_path_nodata = os.path.join(root, ".data.json")
                            file_path_nodata = file_path_nodata.replace('\\', '/')
                            jData = {file_path_nodata: data}
                            
                            try:
                                # Vérifier si le répertoire parent existe
                                parent_dir_path = os.path.dirname(file_path_nodata)
                                if os.path.exists(parent_dir_path):
                                    with open(file_path_nodata, 'w', encoding='utf-8') as file:
                                        json.dump(data, file, ensure_ascii=False, indent=4)
                                    
                                    print(f"{parent_dir}-NEW-4658h : Fichier {file_path_nodata} sauvegardé avec succès")
                                    
                                    # Puis ajouter à la liste de dossiers
                                    dossier_list.append(jData)
                                    record_added = True
                                
                            except Exception as e:
                                print(f"{parent_dir}ERR-4658e Erreur lors de la sauvegarde du fichier {file_path_nodata}: {str(e)}")
                                # Ne pas ajouter à la liste si la sauvegarde a échoué
                                
                    except Exception as e:
                        print(f"{parent_dir}ERR-4658f : Erreur lors de la création d'un nouvel enregistrement: {str(e)}")
        
        print(f"NEW-4658j: Nombre de dossiers traités: {len(dossier_list)}")
        return jsonify(dossier_list), 200
    
    except Exception as e:
        print(f"ERR-4658k : Erreur dans read_annonces_json: {str(e)}")
        import traceback
        traceback.print_exc()
        # En cas d'erreur, retourner un tableau vide plutôt qu'une erreur 500
        return jsonify([]), 200

def load_crit_annonces(excluedFile):
    try:
        #file="excluded_annonces.json"
        config_path = os.path.join(GetRoot(), excluedFile)
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as file:
                return json.load(file)
     
    except Exception as e:
        print(f"Cyr_error_156 An error occurred while loading excluded annonces: {e}")
        return jsonify({"status": "error", "message": "156>"+str(e)}), 500

@cy_routes.route('/save_excluded_annonces', methods=['POST'])
def save_excluded_annonces():
    try:
        data = request.get_json()
        excluded_annonces = data.get('excluded_annonces', [])
        
        dirfilter=GetDirFilter()
        config_path = os.path.join(dirfilter, "excluded_annonces.json")
        with open(config_path, 'w', encoding='utf-8') as config_file:
            json.dump(excluded_annonces, config_file, ensure_ascii=False, indent=4)
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_171 An error occurred while saving excluded annonces: {e}")
        return jsonify({"status": "error", "message": "171>"+str(e)}), 500

'''save config columns'''
@cy_routes.route('/save_config_col', methods=['POST'])
def save_config_col():
    try:
        data = request.get_json()
        serialized_columns = data.get('columns')
        tab_active = data.get('tabActive')
        dirfilter=GetDirFilter()
        # Save the serialized columns to a file or database
        config_path = os.path.join(dirfilter, f"{tab_active}__colums")+ ".json"
        with open(config_path, 'w', encoding='utf-8') as config_file:
            json.dump(serialized_columns, config_file, ensure_ascii=False, indent=4)
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_189 An error occurred while saving config columns: {e}")
        return jsonify({"status": "error", "message": "189>"+str(e)}), 500

# ...existing code...

@cy_routes.route('/open_url', methods=['POST'])
def open_url():
    try:
        data = request.get_json()
        url = data.get('url')
        #print ("##3-------------------------------",url)
        if url:
            # Logic to open the URL
            os.system(f'start {url}')
            return jsonify({"status": "success"}), 200
        else:
            return jsonify({"status": "error", "message": "URL not provided"}), 400
    except Exception as e:
        print(f"Cyr_error_207 An error occurred while opening URL: {e}")
        return jsonify({"status": "error", "message": "207>"+str(e)}), 500

# ...existing code...

@cy_routes.route('/file_exists', methods=['POST'])
def file_exists():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        
        if file_path and os.path.exists(file_path):
            return jsonify({"exists": True}), 200
        else:
            return jsonify({"exists": False}), 200
    except Exception as e:
        print(f"Cyr_error_223 An error occurred while checking file existence: {e}")
        return jsonify({"status": "error", "message": "223>"+str(e)}), 500


@cy_routes.route('/read_filters_json', methods=['POST'])
def read_filters_json():
    try:
        
        data = request.get_json()
        tab_active = data.get('tabActive')
        if not tab_active:
            return {}
        dirfilter=GetDirFilter()
        print(f"dbg-545-loading filters from {dirfilter}")
        file_path = os.path.join(dirfilter, tab_active + "_filter") + ".json"
        file_path = file_path.replace('\\', '/')  # Normalize path
        #print(f"##01-loading filters from {file_path}")
        if not os.path.exists(file_path):
            return jsonify({})
        with open(file_path, 'r', encoding='utf-8') as file:
            filters = json.load(file)
            #print("##02#",filters)
            return jsonify(filters)  # Return dictionary directly
    except Exception as e:
        print(f"Cyr_error_244 An unexpected error occurred while reading filter values: {e}")
        return jsonify({"status": "error", "message": "244>"+str(e)}), 500

# ...existing code...

@cy_routes.route('/save_annonces_json', methods=['POST'])
def save_annonces_json():
    try:
        data = request.get_json()
        for item in data:
            for file_path, content in item.items():
                file_path = file_path.replace('\\', '/')  # Normalize path
                
                # Vérifier si le répertoire parent existe
                parent_dir_path = os.path.dirname(file_path)
                if not os.path.exists(parent_dir_path):
                    print(f"Création du répertoire manquant: {parent_dir_path}")
                    os.makedirs(parent_dir_path, exist_ok=True)
                    
                with open(file_path, 'w', encoding='utf-8') as file:
                    json.dump(content, file, ensure_ascii=False, indent=4)
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_260 An unexpected error occurred while saving data: {e}")
        return jsonify({"status": "error", "message": "260>"+str(e)}), 500

# ...existing code...

@cy_routes.route('/save_filters_json', methods=['POST'])
def save_filters_json():
    try:
        data = request.get_json()
        filters = data.get('filters')
        tab_active = data.get('tabActive')
        
        # Vérifier si tab_active est vide ou null
        if not tab_active:
            print("DEBUG: tab_active est vide, impossible de sauvegarder les filtres")
            return jsonify({"status": "warning", "message": "Onglet actif non spécifié, filtres non sauvegardés"}), 200
            
        dirfilter = GetDirFilter()
        file_path = os.path.join(dirfilter, tab_active + "_filter") + ".json"
        file_path = file_path.replace('\\', '/')  # Normalize path
        print(f"DEBUG: Sauvegarde des filtres dans {file_path}")
        
        # S'assurer que le répertoire existe
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        with open(file_path, 'w', encoding='utf-8') as file:
            json.dump(filters, file, ensure_ascii=False, indent=4)
            
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_277 An unexpected error occurred while saving filter values: {e}")
        return jsonify({"status": "error", "message": "277>"+str(e)}), 500

# ...existing code...

@cy_routes.route('/load_config_col', methods=['POST'])
def load_config_col():
    try:
        data = request.get_json()
        tab_active = data.get('tabActive')
        dirfilter= GetDirFilter()
        file_path = os.path.join(dirfilter, tab_active + "_colums") + ".json"
        file_path = file_path.replace('\\', '/')  # Normalize path
        if not os.path.exists(file_path):
            return jsonify([])  # Return empty list if file does not exist
        with open(file_path, 'r', encoding='utf-8') as file:
            conf = json.load(file)
            return jsonify(conf)  # Return JSON response
    except Exception as e:
        print(f"Cyr_error_295 An unexpected error occurred while reading columns config: {e}")
        return jsonify([])

# ...existing code...

@cy_routes.route('/read_csv_file', methods=['POST'])
def read_csv_file():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        file_path = os.path.join(os.getenv("SUIVI_DIR"), file_path)
        file_path = file_path.replace('\\', '/')  # Normalize path
        #print(f"file path: {file_path}")
        if not os.path.exists(file_path):
            return jsonify([])

        with open(file_path, 'r', encoding='ISO-8859-1') as csvfile:
            csvreader = csv.DictReader(csvfile, delimiter=';')
            data = [row for row in csvreader]
            return jsonify(data)
    except Exception as e:
        print(f"Cyr_error_316 An unexpected error occurred while reading CSV file: {e}")
        return jsonify([])


@cy_routes.route('/save_csv_file', methods=['POST'])
def save_csv_file():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        csv_data = data.get('data')
        file_path = os.path.join(os.getenv("SUIVI_DIR"), file_path)
        file_path = file_path.replace('\\', '/')  # Normalize path
        #print(f"Saving CSV to {file_path}")
        if not os.path.exists(file_path):
            return jsonify({"status": "error", "message": "File path does not exist"}), 400

        with open(file_path, 'w', newline='', encoding='ISO-8859-1') as csvfile:
            fieldnames = csv_data[0].keys()
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames, delimiter=';')
            writer.writeheader()
            writer.writerows(csv_data)
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_339 An unexpected error occurred while saving CSV file: {e}")
        return jsonify({"status": "error", "message": "339>"+str(e)}), 500

# ...existing code...

@cy_routes.route('/open_parent_directory', methods=['POST'])
def open_parent_directory():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        parent_directory = os.path.dirname(file_path)
        parent_directory = parent_directory.replace('\\', '/')  # Normalize path
        if platform.system() == 'Windows':
            os.startfile(parent_directory)
        elif platform.system() == 'Darwin':  # macOS
            subprocess.run(['open', parent_directory])
        else:  # Linux
            subprocess.run(['xdg-open', parent_directory])
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_359 opening parent directory: {e}")
        return jsonify({"status": "error", "message": "359>"+str(e)}), 500

# ...existing code...


@cy_routes.route('/convert_cv', methods=['POST'])
def convert_cv():
    #print ("convertir le cv docx en cv pdf")
    data = request.get_json()
    file_path= data.get('repertoire_annonces')
    num_dossier = data.get('num_dossier')
    filename = secure_filename(f"{num_dossier}_CyrilSauret.docx")
    target_path = os.path.join(file_path, filename)
    target_path_pdf= os.path.join(file_path, filename.replace('.docx', '.pdf'))
    if os.path.exists(target_path_pdf):
        os.remove(target_path_pdf)
        #print("-->04 pdf removed", target_path_pdf)
    

@cy_routes.route('/upload_doc', methods=['POST'])
def select_cv():
    try:
        #print("##0---")
        file = request.files.get('file_path')
        dossier_number = request.form.get('num_dossier')
        target_directory = request.form.get('repertoire_annonce')
        prefix = request.form.get('prefix')
        #print("##2-------------------------------", dossier_number, target_directory)
        
        if not dossier_number or not target_directory or not file:
            return jsonify({"status": "error", "message": "Missing parameters"}), 400 
        
        # Vérifier la taille du fichier avec app_config
        if file.content_length > app_config.max_file_size:
            return jsonify({"status": "error", "message": f"File too large. Maximum size: {app_config.max_file_size // 1024 // 1024}MB"}), 400
        
        filename = secure_filename(f"{dossier_number}_{prefix}_CyrilSauret.docx")
        
        # Valider le nom de fichier avec SecurityValidator
        try:
            target_path = app_config.get_upload_path(filename)
            # Utiliser le répertoire cible au lieu du dossier uploads par défaut
            target_path = os.path.join(target_directory, filename)
            target_path = target_path.replace('\\', '/') 
        except ValueError as e:
            return jsonify({"status": "error", "message": str(e)}), 400
        
        #print("##3-------------------------------", target_path)
        pdf_file_path = target_path.replace('.docx', '.pdf')
        pdf_file_path = pdf_file_path.replace('\\', '/') 
        
        if not os.path.exists(target_path):
            file.save(target_path)
            #print("-->01 docx saved : ", target_path)
        else:
            os.remove(target_path)
            #print("-->02 docx removed", target_path)
            file.save(target_path)
            #print("-->03 docx saved", target_path)
        convert_to_pdf(target_path,pdf_file_path)
        #print("##3a-----------------------",pdf_file_path)
        
        return jsonify({"status": "success", "message": f"File saved as {filename} in {target_directory}"}), 200
    except Exception as e:
        print(f"Cyr_error_411 An error occurred when duplicate file: {e}")
        return jsonify({"status": "error", "message": "441>"+str(e)}), 500 

# ...existing code...


def convert_to_pdf(target_path,pdf_file_path):
    
    if os.path.exists(pdf_file_path):# Delete the existing file
        os.remove(pdf_file_path)
        #print("-->04 pdf removed", pdf_file_path)
    pythoncom.CoInitialize()
    
    try:
        convert(target_path, pdf_file_path)
        #print("-->05 docx converted to pdf", pdf_file_path)
    finally:
        # Uninitialize COM library
        pythoncom.CoUninitialize()
            

def define_default_data():
    return {
        "id": "",
        "description": "?",
        "etat": "Auto",
        "entreprise": "?",
        "categorie": "",
        "Date": "",
        "todo": "?",
        "todoList": "",
        "action": "",
        "tel": "",
        "contact": "",
        "url": "",
        "Commentaire": "",
        "type": "AN",
        "type_question": "pdf",
        "title":"",
       
    }

@cy_routes.route('/save_announcement', methods=['POST'])
def save_announcement():
    try:
        data = request.get_json()
        num_dossier = data.get('contentNum')
        content = data.get('content')
        url = data.get('url')
        sufix = data.get('sufix')

        if not num_dossier or not content or not url:
            print("dbg4456 -------------------------------", num_dossier, content, url)
            return jsonify({"status": "error", "message": "Missing parameters"}), 400

        directory_path = os.path.join(GetRoot(), num_dossier)
        if not os.path.exists(directory_path):
            #print("dbg-4456 creating directory", directory_path)
            os.makedirs(directory_path)
       
        docx_file_path = os.path.join(directory_path, f"{num_dossier}{sufix}.docx")
        pdf_file_path = os.path.join(directory_path, f"{num_dossier}{sufix}.pdf")
        
        """ if os.path.exists(pdf_file_path):
            return jsonify({"status": "error", "message": f"Fichier {pdf_file_path} existe déjà"}), 400 """
        
        pythoncom.CoInitialize()
        # Create DOCX file
        try:
            doc = Document()
            # Ensure URL is properly formatted
            doc.add_paragraph("<-")
            doc.add_paragraph(url)
            doc.add_paragraph("->")
            doc.add_paragraph(content)
            
            doc.save(docx_file_path)
            
            # Convert DOCX to PDF
            print (f"dbg-5434 : Converting {docx_file_path} to {pdf_file_path}")
            convert(docx_file_path, pdf_file_path)
        finally:
            # Uninitialize COM library
            pythoncom.CoUninitialize() 

        return jsonify({"status": "success", "message": f"Announcement saved as {pdf_file_path}"}), 200
    except Exception as e:
        print(f"Cyr_error_492 An error occurred while saving the announcement: {e}")
        return jsonify({"status": "error", "message": "492>"+str(e)}), 500

# ...existing code...

@cy_routes.route('/read_notes', methods=['POST'])
def read_notes():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        
        if not file_path:
            return jsonify({"status": "error", "message": "File path not provided"}), 400
        
        if not os.path.exists(file_path):
            # Create the file if it does not exist
            with open(file_path, 'w', encoding='utf-8') as file:
                json.dump([], file)  # Write an empty JSON array to create the file
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = json.load(file)
        
        return jsonify({"status": "success", "content": content}), 200
    except Exception as e:
        print(f"Cyr_error_516 An error occurred while reading notes: {e}")
        return jsonify({"status": "error", "message": "516>"+str(e)}), 500

# ...existing code...

@cy_routes.route('/save_notes', methods=['POST'])
def save_notes():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        content = data.get('content')
        if not file_path or content is None:
            return jsonify({"status": "error", "message": "Missing parameters"}), 400
        
        # Vérifier si le répertoire parent existe
        parent_dir_path = os.path.dirname(file_path)
        if not os.path.exists(parent_dir_path):
            print(f"Création du répertoire manquant: {parent_dir_path}")
            #os.makedirs(parent_dir_path, exist_ok=True)
        
            with open(file_path, 'w', encoding='utf-8') as file:
                json.dump(content, file, ensure_ascii=False, indent=4)
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_536 An error occurred while saving notes: {e}")
        return jsonify({"status": "error", "message": "536>"+str(e)}), 500



@cy_routes.route('/charger_cols_file', methods=['GET'])
def load_conf_cols():
    try:
        dir=GetRoot()
        filepath = os.path.join(dir, ".cols")
        filepath = filepath.replace('\\', '/')
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as file:
                content = json.load(file)
            print("dbg12391 :fichier conf",filepath)
            #print("dbg12391 :content ",content)   
            return jsonify(content)
        else:
           return jsonify({})
    except Exception as e:
        logger.error(f"Error loading columns configuration: {str(e)}")
        return jsonify({"error": str(e)}), 500

@cy_routes.route('/load_conf_tabs', methods=['GET'])
def load_conf_tabs():
    try:
        dir = GetRoot()
        # Log the root directory for debugging
        print(f"DEBUG: GetRoot() returned: {dir}")
        
        # Verify that dir is a valid string
        if dir is None or not isinstance(dir, str):
            logger.error(f"GetRoot() returned an invalid value: {dir}")
            return jsonify({"error": "Invalid root path"}), 500
            
        filepath = os.path.join(dir, ".conf")
        filepath = filepath.replace('\\', '/')
        
        # Log the full path that's being checked
        print(f"DEBUG: Checking for file at: {filepath}")
        
        if os.path.exists(filepath):
            print(f"DEBUG: File exists at {filepath}")
            with open(filepath, 'r', encoding='utf-8') as file:
                content = json.load(file)
            print(f"DEBUG: Loaded content: {content}")
            return jsonify(content)
        else:
            # Create a default .conf file if it doesn't exist
            print(f"DEBUG: File does not exist at {filepath}, creating default configuration")
            default_conf = {
                "Tabs": [
                  
                ]
            }
            
            # Ensure the directory exists
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # Write the default configuration
            with open(filepath, 'w', encoding='utf-8') as file:
                json.dump(default_conf, file, ensure_ascii=False, indent=4)
            
            return jsonify(default_conf)
    except Exception as e:
        error_msg = f"Error loading tabs configuration: {str(e)}"
        logger.error(error_msg)
        print(f"DEBUG ERROR: {error_msg}")
        return jsonify({"error": error_msg}), 500




async def load_Instruction_classement():
    try:
        text = ""
        file_name_txt = ".clas"
        
        # Essayer d'abord le répertoire actuel (défini par GetRoot())
        current_root = GetRoot()
        filepath = os.path.join(current_root, file_name_txt)
        filepath = filepath.replace('\\', '/')
        print(f"[DEBUG] Recherche du fichier .clas dans le répertoire courant: {filepath}")
        
        # Si le fichier n'existe pas dans le répertoire actuel, essayer le répertoire racine original
        if not os.path.exists(filepath):
            original_root = os.getenv("ANNONCES_FILE_DIR")
            if original_root:
                filepath = os.path.join(original_root, file_name_txt)
                filepath = filepath.replace('\\', '/')
                print(f"[DEBUG] Fichier .clas non trouvé dans le répertoire courant. Recherche dans le répertoire d'origine: {filepath}")
            
            # Si toujours pas trouvé, on peut chercher dans le répertoire du script
            if not os.path.exists(filepath):
                script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # Remonte au répertoire parent du backend
                filepath = os.path.join(script_dir, file_name_txt)
                filepath = filepath.replace('\\', '/')
                print(f"[DEBUG] Fichier .clas non trouvé dans le répertoire d'origine. Recherche dans le répertoire du script: {filepath}")
        
        # Lire le fichier s'il existe
        if os.path.exists(filepath):
            print(f"[INFO] Fichier .clas trouvé à: {filepath}")
            with open(filepath, 'r', encoding='utf-8') as file:
                text = file.read()
        else:
            print(f"[WARN] Fichier .clas introuvable après recherche dans plusieurs répertoires")
        
        return text

    except Exception as e:
        logger.error(f"Error loading text: {str(e)}")
        print(f"[ERROR] Exception lors du chargement du fichier .clas: {str(e)}")
        return ""

@cy_routes.route('/select_dir', methods=['GET'])
async def SelectDirectory():
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    selected_dir = filedialog.askdirectory()
    root.destroy()
    return selected_dir

@cy_routes.route('/generate_html_index', methods=['POST'])
def generate_html_index():
    try:
        data = request.get_json()
        dossier_list = data.get('dossier_list', [])
        sufix = data.get('sufix')
        # Sort the dossier_list by the 'todo' field
        sorted_dossier_list = sorted(dossier_list , key=lambda x: list(x.values())[0].get('todo', ''))
        index_file = os.path.basename(os.path.dirname(GetRoot()))
        index_path = os.path.join(GetRoot(), index_file+ "_index_.html")
        # Check if the file exists to add table headers only once
        file_exists = os.path.exists(index_path)
        if file_exists:
            os.remove(index_path)
        
        with open(index_path, 'a', encoding='utf-8') as index_file:
            index_file.write("<html><body><table border='1'>")
            index_file.write("<tr>")
            index_file.write("<th>N°</th>")
            index_file.write("<th>Entreprise</th>")
            index_file.write("<th>Description du Poste</th>")
            index_file.write("<th>Date de Réponse</th>")
            index_file.write("<th>Todo</th>")
            index_file.write("<th>Commentaire</th>")
            index_file.write("</tr>")
            
            for item in sorted_dossier_list:
                for file_path, data in item.items():
                    index_file.write("<tr>")
                    dossier = data.get('dossier')
                    categorie = data.get('categorie')
                    desc=f"{data.get('categorie', '')} - {data.get('description', '')}"
                    url = data.get('url')
                    index_file.write(f"<td><a href='{dossier}/{dossier}{sufix}.pdf'>{dossier}</a></td>")
                    index_file.write(f"<td>{data.get('entreprise', '')}</td>")
                    index_file.write(f"<td><a href='{url}'>{desc}</a></td>")
                    if categorie == 'Profile':
                        index_file.write(f"<td>{data.get('Date', '')}</td>")
                    elif categorie == 'Annonce':
                        index_file.write(f"<td>{data.get('date_rep', '')}</td>")
                    else:
                        index_file.write(f"<td>{data.get('Date_from', '')}</td>")
                    index_file.write(f"<td>{data.get('todo', '')}</td>")
                    index_file.write(f"<td>{data.get('Commentaire', '')}</td>")
                    index_file.write("</tr>")
            
            index_file.write("</table></body></html>")
        
        # Open the index_path in the file explorer
        os.startfile(index_path)
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Error generating HTML index: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500

# ...existing code...

import os
import shutil

@cy_routes.route('/move_and_rename_directory', methods=['POST'])
def move_and_rename_directory():
    data = request.get_json()
    src_dir = data.get('src_dir')
    dest_dir = data.get('dest_dir') 
    old_prefix = data.get('old_prefix')
    new_prefix = data.get('new_prefix')     
    """
    Déplace un répertoire entier avec tous ses fichiers et remplace les fichiers
    commençant par 'XXXX' par un nouveau préfixe.

    :param src_dir: Chemin du répertoire source à déplacer.
    :param dest_dir: Chemin du répertoire de destination.
    :param new_prefix: Nouveau préfixe pour remplacer 'XXXX' dans les noms de fichiers.
    """
    try:
        if not os.path.exists(src_dir):
            print(f"Le répertoire source '{src_dir}' n'existe pas.")
            return

        if not os.path.exists(dest_dir):
            os.makedirs(dest_dir)
            print(f"Le répertoire de destination '{dest_dir}' a été créé.")

        # Parcourir tous les fichiers et sous-dossiers dans le répertoire source
        for root, dirs, files in os.walk(src_dir):
            # Calculer le chemin relatif pour recréer la structure dans le répertoire de destination
            relative_path = os.path.relpath(root, src_dir)
            target_path = os.path.join(dest_dir, relative_path)

            # Créer les sous-dossiers dans le répertoire de destination
            if not os.path.exists(target_path):
                os.makedirs(target_path)

            # Parcourir les fichiers dans le répertoire actuel
            for file_name in files:
                src_file_path = os.path.join(root, file_name)

                # Renommer les fichiers commençant par 'XXXX'
                if file_name.startswith(old_prefix):
                    new_file_name = file_name.replace(old_prefix, new_prefix, 1)
                else:
                    new_file_name = file_name

                dest_file_path = os.path.join(target_path, new_file_name)

                # Déplacer le fichier
                shutil.move(src_file_path, dest_file_path)
                print(f"Fichier déplacé : {src_file_path} -> {dest_file_path}")

        # Supprimer le répertoire source après le déplacement
        shutil.rmtree(src_dir)
        print(f"Répertoire source '{src_dir}' supprimé après déplacement.")

    except Exception as e:
        print(f"Erreur lors du déplacement du répertoire : {e}")


@cy_routes.route('/clear_explorer_cache', methods=['POST'])
def clear_explorer_cache():
    """
    Efface le cache de l'explorateur et réinitialise les configurations temporaires.
    """
    try:
        # Récupération des chemins de dossiers
        dir_filter = GetDirFilter()
        dir_state = GetDirState()
        
        files_deleted = 0
        
        # Suppression des fichiers de filtres
        if os.path.exists(dir_filter):
            for file in os.listdir(dir_filter):
                if file.endswith("_filter.json") or file.endswith("_colums.json"):
                    file_path = os.path.join(dir_filter, file)
                    try:
                        os.remove(file_path)
                        files_deleted += 1
                        print(f"Fichier supprimé: {file_path}")
                    except Exception as e:
                        print(f"Erreur lors de la suppression du fichier {file_path}: {e}")
        
        # Suppression des fichiers d'état
        if os.path.exists(dir_state):
            for file in os.listdir(dir_state):
                if file.endswith(".json") and not file == "excluded_annonces.json":
                    file_path = os.path.join(dir_state, file)
                    try:
                        os.remove(file_path)
                        files_deleted += 1
                        print(f"Fichier supprimé: {file_path}")
                    except Exception as e:
                        print(f"Erreur lors de la suppression du fichier {file_path}: {e}")
        
        return jsonify({
            "status": "success", 
            "message": f"Cache effacé avec succès. {files_deleted} fichiers supprimés."
        }), 200
    except Exception as e:
        print(f"Erreur lors de l'effacement du cache: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500