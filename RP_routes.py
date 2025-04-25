from flask import Blueprint, request, jsonify, send_from_directory, render_template_string
import os
import json
import platform
import csv
import subprocess
from werkzeug.utils import secure_filename
from PyPDF2 import PdfFileReader

import tkinter as tk
from tkinter import filedialog
import threading
from docx2pdf import convert
import pythoncom
from fpdf import FPDF
from docx import Document
import requests
import json
from tqdm import tqdm
from PyPDF2 import PdfFileReader, PdfFileWriter
from io import BytesIO
import logging
from logging import DEBUG
import aiofiles

from paths import *
from cookies import *
from JO_analyse_gpt import extract_text_from_pdf
from RQ_001 import get_mistral_answer

from dotenv import load_dotenv

load_dotenv()
routes = Blueprint('routes', __name__)
logging.basicConfig(level=DEBUG, format='%(asctime)s - (name)s - (levelname)s - (message)s')
logger = logging.getLogger(__name__)

# Define excluded directories
EXCLUDED_DIRECTORIES = ["suivi", "pile", "conf"]  # Add your excluded directories here
files_type=[{"suffix":"ANNONCE_SUFFIX","type":"AN"}]

# Load shared constants
def load_constants():
    config_path = os.path.join(os.path.dirname(__file__), 'config', 'constants.json')
    with open(config_path, 'r') as f:
        return json.load(f)

CONSTANTS = load_constants()


#print('Loaded constants:', CONSTANTS)


@routes.route('/get_constants', methods=['GET'])
def get_constants():
    return jsonify(CONSTANTS)


from datetime import datetime

def calculate_delay(data):
    try:
        today = datetime.today()
        date_from = data.get('date_from', '')
        date_rep = data.get('date_rep', '')
        todo = data.get('todo', '')

        if 'refus' in todo.lower():
            return "dead"

        if date_rep:
            try:
                date_rep_c = datetime.strptime(date_rep, '%d-%m-%Y')
                #print("dbg778 => date_from_c = "+date_from_c)
                return (today - date_rep_c).days
            except ValueError:
                print ("err dbg456= " +  ValueError )
                pass
            
        if date_from:
            try:
                date_from_c = datetime.strptime(date_from, '%d-%m-%Y')
                #print("dbg778 => date_from_c = "+date_from_c)
                return (today - date_from_c).days
            except ValueError:
                print ("err dbg457= " +  ValueError )
                pass

       

        return 'N/A'
    except Exception as e:
        print(f"Error calculating delay: {e}")
        return 'N/A'


    


@routes.route('/read_annonces_json', methods=['POST'])
async def read_annonces_json():
    try:
        isDetectNew="O"
        buildAllPaths()
        data = request.get_json()
        excluedFile = data.get('excluded')
        directory_path = GetRoot()
       
        if not os.path.exists(directory_path):
            print("dbg676 -> Root path not exist")
            return []
        #print("dbg675 -> directory_path",directory_path)
        dossier_list = []
        crit_annonces = load_crit_annonces(excluedFile)
        #print(f"RP-0 scan repertoire annonces for--------------------------------")
        #print(f"###-0 ")
        for root, _, files in os.walk(directory_path):
            parent_dir = os.path.basename(root)
            if parent_dir in EXCLUDED_DIRECTORIES:
                #print(f"RP-1 --Skipping directory: {parent_dir}")
                continue
            
            file_doc = parent_dir + CONSTANTS['FILE_NAMES']['ANNONCE_SUFFIX'] + ".pdf"
            #file_doc_Action = parent_dir + CONSTANTS['FILE_NAMES']['ACTION_SUFFIX'] + ".pdf"
            
            #file_doc_new = parent_dir + CONSTANTS['FILE_NAMES']['SOURCE_SUFFIX_NEW'] + ".pdf"
            # résumé gpt
            file_isGptResum = parent_dir + CONSTANTS['FILE_NAMES']['GPT_REQUEST_SUFFIX']
            file_isGptResum_Path1 = os.path.join(root, file_isGptResum)
            file_isGptResum_Path1 = file_isGptResum_Path1.replace('\\', '/')
              # CV 
            file_cv = parent_dir + CONSTANTS['FILE_NAMES']['CV_SUFFIX'] + ".docx"
            file_cv_pdf = parent_dir + CONSTANTS['FILE_NAMES']['CV_SUFFIX'] + ".pdf"
            file_BA_pdf = parent_dir + CONSTANTS['FILE_NAMES']['BA_SUFFIX_NAME'] + ".pdf"
            file_BA_docx = parent_dir + CONSTANTS['FILE_NAMES']['BA_SUFFIX_NAME'] + ".docx"
            
            file_cv_New = parent_dir + CONSTANTS['FILE_NAMES']['CV_SUFFIX_NEW'] + ".docx"
            file_cv_pdf_New = parent_dir + CONSTANTS['FILE_NAMES']['CV_SUFFIX_NEW'] + ".pdf"
            data_json_file = ".data.json" 
             
            #file_cv_Path = os.path.join(root, file_cv.replace('\\', '/'))
            record_added = False
            data = {}
            isCVin="N"
            isBAdocx="N"
            isCVinpdf="N"
            isBAinpdf="N"
            isAction="N"
            isJo="N"
          
            file_path_isJo=""
            isGptResum=""
           
        
            isJoTyp=""
            for filename in files:   
                if filename == file_BA_docx: 
                   isBAdocx="O"          
                if filename  == file_cv or filename == file_cv_New:
                    isCVin="O"
                    #print("###---->BINGO")
                if filename  == file_cv_pdf or filename == file_cv_pdf_New:
                    isCVinpdf="O"
                if filename  == file_BA_pdf:
                    isBAinpdf="O"
                if (filename ==  file_doc):
                    file_path_isJo =  os.path.join(root, file_doc.replace('\\', '/'))
                    
                    isJo="O"
                # if (filename ==  file_doc): #or (filename == file_doc_new)):
                #     isJo="O"   
                #     if (filename == file_doc_new):
                #         file_path_isJo = os.path.join(root, file_doc_new)
                #     else:
                #         isAction="O"
                #         file_path_isJo = os.path.join(root, file_doc_Action)
                if (filename ==  file_isGptResum ):
                    isGptResum="O"
                    file_path_gpt = os.path.join(root, file_isGptResum)
            # for filename in files:
            #     for file_type in files_type:
            #         suf = file_type["suffix"]
            #         suftext=CONSTANTS['FILE_NAMES'][suf]
            #         #print("###-3 suffixe = ",suftext)
                    
            #         if suftext in filename:
            #             file_doc = parent_dir + suftext + ".pdf"
            #             #print("###-3 file_doc a trouver= ", file_doc)
            #             if ((filename == file_doc)):
            #                 isJo = "O"
            #                 isJoType = file_type["type"]
            #                 file_path_isJo = os.path.join(root, file_doc)
                            
                       
            for filename in files:
                file_path = os.path.join(root, filename)
                file_path = file_path.replace('\\', '/')  # Normalize path
                file_path_nodata=os.path.join(root, ".data.json")
                file_path_nodata =file_path_nodata.replace('\\', '/') 
                
                if filename == data_json_file :
                    record_added = True
                    try:
                        with open(file_path, 'r', encoding='utf-8') as file:
                            data = json.load(file)
                            # Check exclusion criteria
                            isExclued = False
                            if crit_annonces:
                                excl = crit_annonces["exclude"]
                                for crit in excl:
                                    for key, values in crit.items():
                                        if data.get(key) in values:
                                            isExclued = True
                                            break
                                    if isExclued:
                                        break
                            else:
                                if not data['etat'] in ["DELETED"]:
                                    isExclued = True
                            if not isExclued:
                                data["dossier"] = parent_dir  # Add parent directory name to data
                                if os.path.exists(file_isGptResum_Path1):
                                    isGptResum = "O"# Ensure to handle the case when the file exists.
                                else:
                                    isGptResum = "N"
                                # block ctrl document
                                data["isJo"] = isJo
                               
                                data["GptSum"] = isGptResum
                                data["CV"] = isCVin
                                data["CVpdf"] = isCVinpdf
                                data["BA"] = isBAdocx
                                data["isAction"] = isAction
                                data["BApdf"] = isBAinpdf
                                
                               
                                if "role" not in data:
                                    data["role"] = "default"
                                data["delay"] = calculate_delay(data)
                                jData = {file_path: data}
                                dossier_list.append(jData)
                                """  with open(file_path, 'w', encoding='utf-8') as file:
                                    json.dump(jData, file, ensure_ascii=False, indent=4) """
                            

                    except json.JSONDecodeError:
                        errordata = {"id": parent_dir, "description": "?", "etat": "invalid JSON"}
                        print(f"Cyr_Error: The file {file_path} contains invalid JSON.")
           
            Piece_exist=False
            
            if not record_added:
              
                thefile=""
                Piece_exist=False
                
                    
                if isDetectNew  =="O":
                    if isJo =="O":
                        thefile= file_path_isJo
                        print ("dbg-1245 file_path_isJo trouvé = ",file_path_isJo)
                        Piece_exist=True
                    elif isGptResum =="O":
                        thefile =file_path_gpt
                        Piece_exist=True
                
                if Piece_exist:
                    
                    try:
                        print ("RP-7245 le fichier main va être traité = ",thefile)
                        #thefile = thefile.replace('\\', '/')
                        texte=extract_text_from_pdf(thefile)
                        infos=texte
                        # Use the value of 'current_instruction' from cookies
                        #current_instruction = get_cookie_value('current_instruction')
                            
                        the_request = await load_Instruction_classement()
                        print (f"dbg 6789 {the_request}")
                        if not the_request or the_request.strip() == "":
                             print("Error: the_request is invalid or empty.")
                             #return jsonify({"status": "error", "message": "Invalid instruction request"}), 400
                             infos=texte
                        else:
                             print("RP-2158", the_request)  
                             role="analyse le texte suivant et réponds à cette question, peux tu renvoyer les informations sous forme de données json, les champs son définie dans la question entre [ et ]"
                             infos = get_mistral_answer(the_request, role, texte)
                        print("RP-999 infos = ", infos)
                        if infos:
                            try:
                                # Tenter de parser comme JSON
                                parsed_json = json.loads(infos)
                                data["url"] = parsed_json.get("url", "N/A")
                                data["Date"] = parsed_json.get("Date", "N/A") 
                                data["entreprise"] = parsed_json.get("entreprise", "N/A")
                                data["description"] = parsed_json.get("poste", "N/A")
                                data["Lieu"] = parsed_json.get("lieu", "N/A")
                            except json.JSONDecodeError:
                                # La réponse n'est pas du JSON valide
                                print("RP-1000 : Réponse non JSON, tentative d'extraction des infos du texte")
                                # Extraction basique (peut être améliorée)
                                try:
                                    # Tenter de trouver des données structurées dans la réponse texte
                                    import re
                                    
                                    # Chercher un objet JSON dans la réponse
                                    json_match = re.search(r'(\{.*\})', infos, re.DOTALL)
                                    if json_match:
                                        try:
                                            extracted_json = json.loads(json_match.group(1))
                                            data["url"] = extracted_json.get("url", "N/A")
                                            data["Date"] = extracted_json.get("Date", "N/A")
                                            data["entreprise"] = extracted_json.get("entreprise", "N/A")
                                            data["description"] = extracted_json.get("poste", "N/A")
                                            data["Lieu"] = extracted_json.get("lieu", "N/A")
                                        except:
                                            # Utiliser le texte brut
                                            data["url"] = "N/A"
                                            data["Date"] = "N/A" 
                                            data["entreprise"] = "N/A"
                                            data["description"] = "Pas d'infos"
                                            data["Lieu"] = "N/A"
                                    else:
                                        # Utiliser le texte brut
                                        data["url"] = "N/A"
                                        data["Date"] = "N/A" 
                                        data["entreprise"] = "N/A"
                                        data["description"] = "Pas d'infos"
                                        data["Lieu"] = "N/A"
                                except Exception as extraction_error:
                                    print(f"RP-1001 : Erreur lors de l'extraction: {str(extraction_error)}")
                                    # Fallback en cas d'échec total
                                    data["url"] = "N/A"
                                    data["Date"] = "N/A"
                                    data["entreprise"] = "N/A"  
                                    data["description"] = "Erreur lors du traitement"
                                    data["Lieu"] = "N/A"
                            print ("RP-999 infos = ",infos)    
                            data["dossier"] = parent_dir    
                            data["isJo"] = isJo
                            data["isAction"] = isAction
                            data["GptSum"] = isGptResum
                            data["CV"] = isCVin
                            data["CVpdf"] = isCVinpdf
                                # block info piece         
                            data["etat"] = "New"
                            jData = {file_path_nodata:data}  
                                
                            dossier_list.append(jData)
                            record_added = True
                            #file_path_nodata = file_path_nodata.replace('\\', '/')  # Normalize path
                            with open(file_path_nodata, 'w', encoding='utf-8') as file:
                                    json.dump(data, file, ensure_ascii=False, indent=4)
                    except Exception as e:   
                        print(f"Cyr_Error 14578: An error occurred while trying to retrieve information from {thefile}: {str(e)}")
                        return []
        return dossier_list 
    except Exception as e:
        print(f"Cyr_error_145 An unexpected error occurred while reading annonces: {e}")
        return []

def load_crit_annonces(excluedFile):
    try:
        #file="excluded_annonces.json"
        config_path = os.path.join(GetDirState(), excluedFile)
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as file:
                return json.load(file)
     
    except Exception as e:
        print(f"Cyr_error_156 An error occurred while loading excluded annonces: {e}")
        return jsonify({"status": "error", "message": "156>"+str(e)}), 500

@routes.route('/save_excluded_annonces', methods=['POST'])
def save_excluded_annonces():
    try:
        data = request.get_json()
        excluded_annonces = data.get('excluded_annonces', [])
        
        dirfilter=GetDirFilter()
        config_path = os.path.join(dirfilter, "excluded_annonces.json")
        with open(config_path, 'w', encoding='utf-8') as config_file:
            json.dump(excluded_annonces, config_file, ensure_ascii=False, indent=4)
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_171 An error occurred while saving excluded annonces: {e}")
        return jsonify({"status": "error", "message": "171>"+str(e)}), 500

'''save config columns'''
@routes.route('/save_config_col', methods=['POST'])
def save_config_col():
    try:
        data = request.get_json()
        serialized_columns = data.get('columns')
        tab_active = data.get('tabActive')
        dirfilter=GetDirFilter()
        # Save the serialized columns to a file or database
        config_path = os.path.join(dirfilter, f"{tab_active}__colums")+ ".json"
        with open(config_path, 'w', encoding='utf-8') as config_file:
            json.dump(serialized_columns, config_file, ensure_ascii=False, indent=4)
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_189 An error occurred while saving config columns: {e}")
        return jsonify({"status": "error", "message": "189>"+str(e)}), 500

# ...existing code...

@routes.route('/open_url', methods=['POST'])
def open_url():
    try:
        data = request.get_json()
        url = data.get('url')
        #print ("##3-------------------------------",url)
        if url:
            # Logic to open the URL
            os.system(f'start {url}')
            return jsonify({"status": "success"}), 200
        else:
            return jsonify({"status": "error", "message": "URL not provided"}), 400
    except Exception as e:
        print(f"Cyr_error_207 An error occurred while opening URL: {e}")
        return jsonify({"status": "error", "message": "207>"+str(e)}), 500

# ...existing code...

@routes.route('/file_exists', methods=['POST'])
def file_exists():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        
        if file_path and os.path.exists(file_path):
            return jsonify({"exists": True}), 200
        else:
            return jsonify({"exists": False}), 200
    except Exception as e:
        print(f"Cyr_error_223 An error occurred while checking file existence: {e}")
        return jsonify({"status": "error", "message": "223>"+str(e)}), 500


@routes.route('/read_filters_json', methods=['POST'])
def read_filters_json():
    try:
        
        data = request.get_json()
        tab_active = data.get('tabActive')
        
        dirfilter=GetDirFilter()
        
        file_path = os.path.join(dirfilter, tab_active + "_filter") + ".json"
        file_path = file_path.replace('\\', '/')  # Normalize path
        #print(f"##01-loading filters from {file_path}")
        if not os.path.exists(file_path):
            return jsonify({})
        with open(file_path, 'r', encoding='utf-8') as file:
            filters = json.load(file)
            #print("##02#",filters)
            return jsonify(filters)  # Return dictionary directly
    except Exception as e:
        print(f"Cyr_error_244 An unexpected error occurred while reading filter values: {e}")
        return jsonify({"status": "error", "message": "244>"+str(e)}), 500

# ...existing code...

@routes.route('/save_annonces_json', methods=['POST'])
def save_annonces_json():
    try:
        data = request.get_json()
        for item in data:
            for file_path, content in item.items():
                file_path = file_path.replace('\\', '/')  # Normalize path
                with open(file_path, 'w', encoding='utf-8') as file:
                    json.dump(content, file, ensure_ascii=False, indent=4)
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_260 An unexpected error occurred while saving data: {e}")
        return jsonify({"status": "error", "message": "260>"+str(e)}), 500

# ...existing code...

@routes.route('/save_filters_json', methods=['POST'])
def save_filters_json():
    try:
        data = request.get_json()
        filters = data.get('filters')
        tab_active = data.get('tabActive')
        dirfilter = GetDirFilter()
        file_path = os.path.join(dirfilter, tab_active + "_filter") + ".json"
        file_path = file_path.replace('\\', '/')  # Normalize path
        #print(f"##-9998-saving filters to {file_path}")    
        with open(file_path, 'w', encoding='utf-8') as file:
            json.dump(filters, file, ensure_ascii=False, indent=4)
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_277 An unexpected error occurred while saving filter values: {e}")
        return jsonify({"status": "error", "message": "277>"+str(e)}), 500

# ...existing code...

@routes.route('/load_config_col', methods=['POST'])
def load_config_col():
    try:
        data = request.get_json()
        tab_active = data.get('tabActive')
        dirfilter= GetDirFilter()
        file_path = os.path.join(dirfilter, tab_active + "_colums") + ".json"
        file_path = file_path.replace('\\', '/')  # Normalize path
        if not os.path.exists(file_path):
            return jsonify([])  # Return empty list if file does not exist
        with open(file_path, 'r', encoding='utf-8') as file:
            conf = json.load(file)
            return jsonify(conf)  # Return JSON response
    except Exception as e:
        print(f"Cyr_error_295 An unexpected error occurred while reading columns config: {e}")
        return jsonify([])

# ...existing code...

@routes.route('/read_csv_file', methods=['POST'])
def read_csv_file():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        file_path = os.path.join(os.getenv("SUIVI_DIR"), file_path)
        file_path = file_path.replace('\\', '/')  # Normalize path
        #print(f"file path: {file_path}")
        if not os.path.exists(file_path):
            return jsonify([])

        with open(file_path, 'r', encoding='ISO-8859-1') as csvfile:
            csvreader = csv.DictReader(csvfile, delimiter=';')
            data = [row for row in csvreader]
            return jsonify(data)
    except Exception as e:
        print(f"Cyr_error_316 An unexpected error occurred while reading CSV file: {e}")
        return jsonify([])


@routes.route('/save_csv_file', methods=['POST'])
def save_csv_file():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        csv_data = data.get('data')
        file_path = os.path.join(os.getenv("SUIVI_DIR"), file_path)
        file_path = file_path.replace('\\', '/')  # Normalize path
        #print(f"Saving CSV to {file_path}")
        if not os.path.exists(file_path):
            return jsonify({"status": "error", "message": "File path does not exist"}), 400

        with open(file_path, 'w', newline='', encoding='ISO-8859-1') as csvfile:
            fieldnames = csv_data[0].keys()
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames, delimiter=';')
            writer.writeheader()
            writer.writerows(csv_data)
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_339 An unexpected error occurred while saving CSV file: {e}")
        return jsonify({"status": "error", "message": "339>"+str(e)}), 500

# ...existing code...

@routes.route('/open_parent_directory', methods=['POST'])
def open_parent_directory():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        parent_directory = os.path.dirname(file_path)
        parent_directory = parent_directory.replace('\\', '/')  # Normalize path
        if platform.system() == 'Windows':
            os.startfile(parent_directory)
        elif platform.system() == 'Darwin':  # macOS
            subprocess.run(['open', parent_directory])
        else:  # Linux
            subprocess.run(['xdg-open', parent_directory])
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_359 opening parent directory: {e}")
        return jsonify({"status": "error", "message": "359>"+str(e)}), 500

# ...existing code...


@routes.route('/convert_cv', methods=['POST'])
def convert_cv():
    #print ("convertir le cv docx en cv pdf")
    data = request.get_json()
    file_path= data.get('repertoire_annonces')
    num_dossier = data.get('num_dossier')
    filename = secure_filename(f"{num_dossier}_CyrilSauret.docx")
    target_path = os.path.join(file_path, filename)
    target_path_pdf= os.path.join(file_path, filename.replace('.docx', '.pdf'))
    if os.path.exists(target_path_pdf):
        os.remove(target_path_pdf)
        #print("-->04 pdf removed", target_path_pdf)
    

@routes.route('/share_cv', methods=['POST'])
def select_cv():
    try:
        #print("##0---")
        file = request.files.get('file_path')
        dossier_number = request.form.get('num_dossier')
        target_directory = request.form.get('repertoire_annonce')
        prefix= request.form.get('prefix')
        #print("##2-------------------------------", dossier_number, target_directory)
        
        if not dossier_number or not target_directory or not file:
            return jsonify({"status": "error", "message": "Missing parameters"}), 400 
        
        filename = secure_filename(f"{dossier_number}_{prefix}_CyrilSauret.docx")
        target_path = os.path.join(target_directory, filename)
        target_path = target_path.replace('\\', '/') 
        #print("##3-------------------------------", target_path)
        pdf_file_path = target_path.replace('.docx', '.pdf')
        pdf_file_path = pdf_file_path.replace('\\', '/') 
        
        if not os.path.exists(target_path):
            file.save(target_path)
            #print("-->01 docx saved : ", target_path)
        else:
            os.remove(target_path)
            #print("-->02 docx removed", target_path)
            file.save(target_path)
            #print("-->03 docx saved", target_path)
        convert_to_pdf(target_path,pdf_file_path)
        #print("##3a-----------------------",pdf_file_path)
        
        return jsonify({"status": "success", "message": f"File saved as {filename} in {target_directory}"}), 200
    except Exception as e:
        print(f"Cyr_error_411 An error occurred when duplicate file: {e}")
        return jsonify({"status": "error", "message": "441>"+str(e)}), 500 

# ...existing code...


def convert_to_pdf(target_path,pdf_file_path):
    
    if os.path.exists(pdf_file_path):# Delete the existing file
        os.remove(pdf_file_path)
        #print("-->04 pdf removed", pdf_file_path)
    pythoncom.CoInitialize()
    
    try:
        convert(target_path, pdf_file_path)
        #print("-->05 docx converted to pdf", pdf_file_path)
    finally:
        # Uninitialize COM library
        pythoncom.CoUninitialize()
            

def define_default_data():
    return {
        "id": "",
        "description": "?",
        "etat": "Auto",
        "entreprise": "?",
        "categorie": "",
        "Date": "",
        "todo": "?",
        "todoList": "",
        "action": "",
        "tel": "",
        "contact": "",
        "url": "",
        "Commentaire": "",
        "type": "AN",
        "type_question": "pdf",
        "title":"",
       
    }

@routes.route('/save_announcement', methods=['POST'])
def save_announcement():
    try:
        data = request.get_json()
        num_dossier = data.get('contentNum')
        content = data.get('content')
        url = data.get('url')
        sufix = data.get('sufix')

        if not num_dossier or not content or not url:
            return jsonify({"status": "error", "message": "Missing parameters"}), 400

        directory_path = os.path.join(GetRoot(), num_dossier)
        if not os.path.exists(directory_path):
            os.makedirs(directory_path)
       
        docx_file_path = os.path.join(directory_path, f"{num_dossier}{sufix}.docx")
        pdf_file_path = os.path.join(directory_path, f"{num_dossier}{sufix}.pdf")
        
        """ if os.path.exists(pdf_file_path):
            return jsonify({"status": "error", "message": f"Fichier {pdf_file_path} existe déjà"}), 400 """
        
        pythoncom.CoInitialize()
        # Create DOCX file
        try:
            doc = Document()
            # Ensure URL is properly formatted
            doc.add_paragraph("<-")
            doc.add_paragraph(url)
            doc.add_paragraph("->")
            doc.add_paragraph(content)
            
            doc.save(docx_file_path)
            # Convert DOCX to PDF
            convert(docx_file_path, pdf_file_path)
        finally:
            # Uninitialize COM library
            pythoncom.CoUninitialize() 

        return jsonify({"status": "success", "message": f"Announcement saved as {pdf_file_path}"}), 200
    except Exception as e:
        print(f"Cyr_error_492 An error occurred while saving the announcement: {e}")
        return jsonify({"status": "error", "message": "492>"+str(e)}), 500

# ...existing code...

@routes.route('/read_notes', methods=['POST'])
def read_notes():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        
        if not file_path:
            return jsonify({"status": "error", "message": "File path not provided"}), 400
        
        if not os.path.exists(file_path):
            # Create the file if it does not exist
            with open(file_path, 'w', encoding='utf-8') as file:
                json.dump([], file)  # Write an empty JSON array to create the file
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = json.load(file)
        
        return jsonify({"status": "success", "content": content}), 200
    except Exception as e:
        print(f"Cyr_error_516 An error occurred while reading notes: {e}")
        return jsonify({"status": "error", "message": "516>"+str(e)}), 500

# ...existing code...

@routes.route('/save_notes', methods=['POST'])
def save_notes():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        content = data.get('content')
        
        if not file_path or content is None:
            return jsonify({"status": "error", "message": "Missing parameters"}), 400
        
        with open(file_path, 'w', encoding='utf-8') as file:
            json.dump(content, file, ensure_ascii=False, indent=4)
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_536 An error occurred while saving notes: {e}")
        return jsonify({"status": "error", "message": "536>"+str(e)}), 500

# ...existing code...

# @routes.route('/load_reseaux_link', methods=['GET'])
# def load_reseaux_link():
#     try:
#        file_path = os.path.join(os.getenv("RESEAUX_FILE"))
#        file_path = file_path.replace('\\', '/')  # Normalize path
#        if not os.path.exists(file_path):
#            return jsonify([])  # Return empty list if file does not exist
#        with open(file_path, 'r', encoding='utf-8') as file:
#            file = json.load(file)
#            return jsonify(file)  # Return JSON response
#     except Exception as e:
#         print(f"Cyr_error_552 An unexpected error occurred while reading reseaux: {e}")
#         return jsonify([])

""" @routes.route('/save_reseaux_link_update', methods=['POST'])
def save_reseaux_link_update():
    try:
        link_data = request.get_json()
        file_path = os.getenv("RESEAUX_FILE")
        file_path = file_path.replace('\\', '/')  # Normalize path
        #print(f"Cyr_error_561---",file_path)
        if not os.path.exists(file_path):
            return jsonify({"status": "error", "message": "File path does not exist"}), 400

        with open(file_path, 'r', encoding='utf-8') as file:
            links = json.load(file)

        # Update the link in the list
        for link in links:
            if link['name'] == link_data['name']:
                link.update(link_data)
                break

        with open(file_path, 'w', encoding='utf-8') as file:
            json.dump(links, file, ensure_ascii=False, indent=4)

        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Cyr_error_579 An error occurred while saving reseaux link update: {e}")
        return jsonify({"status": "error", "message": "579>"+str(e)}), 500
 """
# ...existing code...

@routes.route('/list-CRQ-files', methods=['GET'])
async def list_CRQ_files():
    list_CRQ = []
    try:
        directory_path = GetDirCRQ('DIR_DRQ_FILE')
        if not os.path.exists(directory_path):
           os.makedirs(directory_path) 
        defaultfile = os.path.join(directory_path, 'default.txt')
        if not os.path.exists(directory_path):
            #logger.info("Directory does not exist, creating: %s", directory_path)
            os.makedirs(directory_path)
            the_request = (
                " -peux tu trouver : l'url référence de l'annonce ['url'], l'url peut aussi se trouver entre <- et ->, "
                "-l'entreprise [entreprise], "
                "-le titre ou l'intiltulé [poste] du poste à pourvoir (ce titre ne doit pas dépasser 20 caractères), "
                "-la localisation ou lieu dans lieux [lieu], "
                "-la date de publication ou d'actualisation [Date]"
            )
            defaultfile = defaultfile.replace('\\', '/')
            #print("dbg1456-----", defaultfile)
            await save_CRQ_text(defaultfile,the_request)
        
        text_files = [f for f in os.listdir(directory_path) if f.endswith('.txt')]
        #for each file in directory_path
        for file in text_files:
            file_path = os.path.join(directory_path, file)
            with open(file_path, 'r', encoding='utf-8') as thefile:
                list_CRQ.append({
                    'fichier': file_path,
                    'name': file,
                    
                })
        
        return jsonify(list_CRQ), 200
    except Exception as e:
        logger.error("Unable to scan directory: %s", str(e))
        return jsonify({"error": f"Unable to scan directory: {str(e)}"}), 500

# ...existing code...

@routes.route('/save-CRQ-text', methods=['POST'])
async def route_save_CRQ_text(): 
    file_name = request.json.get('file_name')
    text_data = request.json.get('text_data')
    #print("dbg897 :fichier name", file_name)
    #print("dbg897a :text_data", text_data)
    
    return  save_CRQ_text(file_name, text_data)

def save_CRQ_text(file_name, text_data):
    try:
       
        if not file_name or not text_data:
            return jsonify({'error': 'Missing file name or text data'}), 400
        #print("dbg897b :sauvegarde en cours ")
        with open(file_name, 'w', encoding='utf-8') as file:
            file.write(text_data)
        
        #logger.debug(f"Text saved successfully as {file_name}")
        return jsonify({'message': 'Text saved successfully'}), 200

    except Exception as e:
        logger.error(f"Error saving text: {str(e)}")
        return jsonify({'error': str(e)}), 500

# ...existing code...



@routes.route('/load-conf-cols', methods=['GET'])
def load_conf_cols():
    
    dir=GetRoot()
    
    filepath = os.path.join(dir, ".cols")
    filepath = filepath.replace('\\', '/')
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as file:
            content = json.load(file)
        print("dbg12391 :fichier conf",filepath)
        print("dbg12391 :content ",content)   
        return content
    else:
        return jsonify({"error": "Configuration file does not exist"}), 404
    



@routes.route('/load-conf-tabs', methods=['GET'])
def load_conf_tabs():
    
    dir=GetRoot()
    
    filepath = os.path.join(dir, ".conf")
    filepath = filepath.replace('\\', '/')
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as file:
            content = json.load(file)
        print("dbg897 :fichier conf",filepath)
        print("dbg897 :content ",content)        
        return content
    else:
        return jsonify({"error": "Configuration file does not exist"}), 404
    

# @routes.route('/load-CRQ-text', methods=['POST'])
# def route_load_CRQ_text():
#     file_name = request.json.get('file_name')
#     dir=request.json.get('dir')
 
#     #print("dbg788 :fichier instructions",file_name)     
#     text=load_CRQ_text(file_name,dir)
#     #print("dbg790 :text ",text)   
#     return jsonify(text)
    
async def load_Instruction_classement():
    try:
        text = ""
        file_name_txt = ".clas"
        filepath = os.path.join(GetRoot(), file_name_txt)  # Updated to call GetRoot() correctly
        filepath = filepath.replace('\\', '/')
        #print("dbg3434 :fichier requete",filepath)
        #print("dbg789 :fichier instructions",filepath)
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as file:
                text = file.read()
        
        return text

    except Exception as e:
        logger.error(f"Error loading text: {str(e)}")
        return ""

@routes.route('/select_dir', methods=['GET'])
async def SelectDirectory():
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    selected_dir = filedialog.askdirectory()
    root.destroy()
    return selected_dir

@routes.route('/generate_html_index', methods=['POST'])
def generate_html_index():
    try:
        data = request.get_json()
        dossier_list = data.get('dossier_list', [])
        sufix = data.get('sufix')
        # Sort the dossier_list by the 'todo' field
        sorted_dossier_list = sorted(dossier_list , key=lambda x: list(x.values())[0].get('todo', ''))
        index_path = os.path.join(GetRoot(), os.getenv("INDEX_DOSSIERS"))
        # Check if the file exists to add table headers only once
        file_exists = os.path.exists(index_path)
        if file_exists:
            os.remove(index_path)
        
        with open(index_path, 'a', encoding='utf-8') as index_file:
            index_file.write("<html><body><table border='1'>")
            index_file.write("<tr>")
            index_file.write("<th>N°</th>")
            index_file.write("<th>Entreprise</th>")
            index_file.write("<th>Description du Poste</th>")
            index_file.write("<th>Date de Réponse</th>")
            index_file.write("<th>Todo</th>")
            index_file.write("<th>Commentaire</th>")
            index_file.write("</tr>")
            
            for item in sorted_dossier_list:
                for file_path, data in item.items():
                    index_file.write("<tr>")
                    dossier = data.get('dossier')
                    categorie = data.get('categorie')
                    desc=f"{data.get('categorie', '')} - {data.get('description', '')}"
                    url = data.get('url')
                    index_file.write(f"<td><a href='{dossier}/{dossier}{sufix}.pdf'>{dossier}</a></td>")
                    index_file.write(f"<td>{data.get('entreprise', '')}</td>")
                    index_file.write(f"<td><a href='{url}'>{desc}</a></td>")
                    if categorie == 'Profile':
                        index_file.write(f"<td>{data.get('Date', '')}</td>")
                    elif categorie == 'Annonce':
                        index_file.write(f"<td>{data.get('date_rep', '')}</td>")
                    else:
                        index_file.write(f"<td>{data.get('Date_from', '')}</td>")
                    index_file.write(f"<td>{data.get('todo', '')}</td>")
                    index_file.write(f"<td>{data.get('Commentaire', '')}</td>")
                    index_file.write("</tr>")
            
            index_file.write("</table></body></html>")
        
        # Open the index_path in the file explorer
        os.startfile(index_path)
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        print(f"Error generating HTML index: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500

# ...existing code...

import os
import shutil

@routes.route('/move_and_rename_directory', methods=['POST'])
def move_and_rename_directory():
    data = request.get_json()
    src_dir = data.get('src_dir')
    dest_dir = data.get('dest_dir') 
    old_prefix = data.get('old_prefix')
    new_prefix = data.get('new_prefix')     
    """
    Déplace un répertoire entier avec tous ses fichiers et remplace les fichiers
    commençant par 'XXXX' par un nouveau préfixe.

    :param src_dir: Chemin du répertoire source à déplacer.
    :param dest_dir: Chemin du répertoire de destination.
    :param new_prefix: Nouveau préfixe pour remplacer 'XXXX' dans les noms de fichiers.
    """
    try:
        if not os.path.exists(src_dir):
            print(f"Le répertoire source '{src_dir}' n'existe pas.")
            return

        if not os.path.exists(dest_dir):
            os.makedirs(dest_dir)
            print(f"Le répertoire de destination '{dest_dir}' a été créé.")

        # Parcourir tous les fichiers et sous-dossiers dans le répertoire source
        for root, dirs, files in os.walk(src_dir):
            # Calculer le chemin relatif pour recréer la structure dans le répertoire de destination
            relative_path = os.path.relpath(root, src_dir)
            target_path = os.path.join(dest_dir, relative_path)

            # Créer les sous-dossiers dans le répertoire de destination
            if not os.path.exists(target_path):
                os.makedirs(target_path)

            # Parcourir les fichiers dans le répertoire actuel
            for file_name in files:
                src_file_path = os.path.join(root, file_name)

                # Renommer les fichiers commençant par 'XXXX'
                if file_name.startswith(old_prefix):
                    new_file_name = file_name.replace(old_prefix, new_prefix, 1)
                else:
                    new_file_name = file_name

                dest_file_path = os.path.join(target_path, new_file_name)

                # Déplacer le fichier
                shutil.move(src_file_path, dest_file_path)
                print(f"Fichier déplacé : {src_file_path} -> {dest_file_path}")

        # Supprimer le répertoire source après le déplacement
        shutil.rmtree(src_dir)
        print(f"Répertoire source '{src_dir}' supprimé après déplacement.")

    except Exception as e:
        print(f"Erreur lors du déplacement du répertoire : {e}")