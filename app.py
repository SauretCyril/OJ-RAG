import warnings
warnings.filterwarnings("ignore", category=UserWarning)

from flask import Flask, render_template, jsonify, request, url_for
import os

from werkzeug.utils import secure_filename
from JO_analyse import *
from JO_analyse_gpt import extract_text_from_url, extract_text_from_pdf, get_answer  # Import the functions from the correct module

''' Procédures'''
from RQ_001 import get_mistral_answer, mistral  # Import the function and Blueprint
# ...existing code...
import torch
import numpy as np
import json
import logging
from docx import Document
from docx2pdf import convert
import pythoncom
from fpdf import FPDF
from dotenv import load_dotenv
from paths import *

# Load environment variables from .env file
load_dotenv()

logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Définir NumpyEncoder avant de l'utiliser
class NumpyEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, np.ndarray):
            return obj.tolist()
        if isinstance(obj, np.float32):
            return float(obj)
        if isinstance(obj, torch.Tensor):
            return obj.detach().cpu().numpy().tolist()
        return super(NumpyEncoder, self).default(obj)

# Configurer Flask pour utiliser le NumpyEncoder
app = Flask(__name__)
app.json_encoder = NumpyEncoder  # Ajouter cette ligne après la création de l'app
from RP_routes import routes 
app.register_blueprint(routes)  # Register the blueprint

from exploreur import exploreur
app.register_blueprint(exploreur)  # Register the blueprint      

#from ST_steal import Steal 
#app.register_blueprint(Steal)  # Register the blueprint

from cookies import cookies
app.register_blueprint(cookies)  # Register the blueprint

#from list_realisations import realizations
#app.register_blueprint(realizations)

#from list_requests import requests
#app.register_blueprint(requests)

# Enregistrer le Blueprint mistral
app.register_blueprint(mistral)  # Register the mistral blueprint

# Ajoutez ce code dans votre app.py après avoir enregistré le Blueprint
""" print("Routes disponibles:")
for rule in app.url_map.iter_rules():
    print(f"{rule} - {rule.endpoint}") """

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'uploads')
app.static_folder = os.path.join(BASE_DIR, 'static')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max-limit

# Assurer que les dossiers nécessaires existent
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

SAVED_TEXT_FOLDER = 'path/to/save/folder'
ALLOWED_EXTENSIONS = {'pdf'}

app.config['SAVED_TEXT_FOLDER'] = SAVED_TEXT_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/alive')
def alive():
    return jsonify({'message': 'Im alive'}), 200

@app.route('/')
def index():
    return render_template('RP_index.html')

@app.route('/get_file_path', methods=['POST'])
def get_file_path():
    dirname = request.json.get('dirName')
    filename = request.json.get('filename')
    print("dbg3214 : dirName", dirname)
    print("dbg3215 : filename", filename)

    if not dirname or not filename:
        return jsonify({'error': 'Invalid parameters'}), 400    
    
    dirnamevalue = GetOneDir(dirname)
    print(f"dbg3216a repertoire du fichier {dirname}: filepath", dirnamevalue)
    
    filenamevalue = os.getenv(filename)
    print(f"dbg3216b nom du fichier = {filenamevalue}")
    
    if not filenamevalue:
        return jsonify({'error': 'Filename not found in environment variables'}), 404
    
    filepath = os.path.join(dirnamevalue, filenamevalue)
    print("dbg3216c : fichier ", filepath)
    
    return jsonify({
        'filePath': filepath,
        'exist': os.path.exists(filepath)
    })

@app.route('/save-path')
def getSavePath():
    return SAVED_TEXT_FOLDER

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            logger.error("Er001.No file part in the request")
            return jsonify({'Er001': 'No file part'}), 400

        file1 = request.files['file']
        
        if file1.filename == '':
            logger.error("Er002.No selected file")
            return jsonify({'Er002': 'No selected file'}), 400

        if not allowed_file(file1.filename):
            logger.error(f"Er003.File type not allowed: {file1.filename}")
            return jsonify({'Er003': 'File type not allowed'}), 400
        
        file1_path = file1.filename[:4]
        if not os.path.exists(app.config['UPLOAD_FOLDER']):
            os.makedirs(app.config['UPLOAD_FOLDER'])

        file = os.path.join(app.config['UPLOAD_FOLDER'],file1.filename)
        file1.save(file)

        return jsonify({
            'path': file,
            'file_dir': file1_path
        })

    except Exception as e:
        logger.error(f"Er004.Upload error: {str(e)}")
        return jsonify({'Er004': str(e)}), 500

@app.route('/job-details')
def job_details():
    return render_template('job_details.html')

@app.route('/get_job_answer', methods=['POST'])
def get_job_answer():
    try:
        file = request.json.get('path')
        RQ = request.json.get('RQ')

        if not file or not RQ:
            logger.error("Er005.Missing job file path or question")
            return jsonify({'Er005': 'Missing job file path or question'}), 400

        text1 = extract_text_from_pdf(file)
        
        if not text1:
            logger.error("Er006.Job text extraction failed")
            return jsonify({'Er006': 'Job text extraction failed'}), 500

        role="En tant qu' expert en analyse d'offres d'emploi dans le domaine informatique (développeur, Analyste ou Testeur logiciel) , analyse le texte suivant et réponds à cette question"
        #answer = get_answer(RQ,role, text1)
        
        answer = get_mistral_answer(RQ,role,text1)
        return jsonify({
            'raw_text': text1,
            'formatted_text': answer
        })

    except Exception as e:
        logger.error(f"Er007.Error: {str(e)}")
        return jsonify({'Er007': str(e)}), 500

@app.route('/save-answer', methods=['POST'])
def save_answer():
    try:
        pythoncom.CoInitialize()  # Initialize COM library
        job_text_data = request.json.get('text_data')
        job_number = request.json.get('number')
        the_path = request.json.get('the_path')
        rq=request.json.get('RQ')
        if the_path == '':
            the_path = GetRoot()

        if not job_text_data or not job_number:
            logger.error(f"Er008.error.Missing job text data or job number: job_text_data={job_text_data}, job_number={job_number}")
            return jsonify({'Er008': 'Missing job text data or job number'}), 400

        if not the_path:
            logger.error("Er009.Received path is None")
            return jsonify({'Er009': 'Missing path'}), 400

        file_name = f"{job_number}_gpt_request"
        file_path = os.path.join(the_path, job_number)
        if not os.path.exists(file_path):
            os.makedirs(file_path)
        
        file_path_docx = os.path.join(file_path, file_name + ".docx")
        file_path_RQ = os.path.join(file_path, file_name + "_RQ.txt")

        doc = format_text_as_word_style(job_text_data, job_number)
        doc.save(file_path_docx)

        pdf_file_path = file_path_docx.replace('.docx', '.pdf')
        convert(file_path_docx, pdf_file_path)
        
        os.remove(file_path_docx)

        save_rq_to_text_file(file_path_RQ, rq)
        
        return jsonify({'dbg009': 'Job text saved successfully', 'pdf_file_path': pdf_file_path})

    except Exception as e:
        logger.error(f"Er009.Error saving job text: {str(e)}")
        return jsonify({'Er009': str(e)}), 500

    finally:
        pythoncom.CoUninitialize()  # Uninitialize COM library

def save_rq_to_text_file(file_path, rq):
    with open(file_path, 'w') as file:
        file.write(rq)

def format_text_as_word_style(job_text, job_number):
    doc = Document()
    doc.add_heading(f"Job Number: {job_number}", level=1)
    
    for line in job_text.split('\n'):
        if line.startswith('- '):
            doc.add_paragraph(line, style='ListBullet')
        else:
            doc.add_paragraph(line)
    
    return doc

@app.route('/extract_features', methods=['POST'])
def extract_features(text):
    cv_features = extract_features("cv_text")
    job_features = extract_features("job_text")
    raw_similarity = cosine_similarity(cv_features, job_features)
    adjusted_similarity = calculate_similarity_score(cv_features, job_features)
        
    raw_similarity_s=f"{raw_similarity:.4f}"
    adjusted_similarity_s=f"{adjusted_similarity:.4f}"
    return jsonify({'Raw_similarity_score': raw_similarity_s, 'Adjusted_similarity_score':adjusted_similarity_s})

@app.route('/get_job_answer_from_url', methods=['POST'])
def get_job_answer_from_url():
    try:
        url = request.json.get('url')
        RQ = request.json.get('RQ')

        if not url or not RQ:
            logger.error("Er014.Missing job file path or question")
            return jsonify({'Er014': 'Missing job file path or question'}), 400

        text1 = extract_text_from_url(url)
        text1 += "<-"+url+"->"
        
        if not text1:
            logger.error("Er0016.Job text extraction failed")
            return jsonify({'Er016': 'Job text extraction failed'}), 500
        role = "En tant qu' expert en analyse d'offres d'emploi dans le domaine informatique (développeur, Analyste ou Testeur logiciel), analyse le texte suivant et réponds à cette question"
        answer = get_mistral_answer(RQ, role, text1)

        return jsonify({
            'raw_text': text1,
            'formatted_text': answer
        })

    except Exception as e:
        logger.error(f"Er018.Error: {str(e)}")
        return jsonify({'Er018': str(e)}), 500

@app.route('/check_dossier_exists', methods=['POST'])
def check_dossier_exist():
    try:
        directory_path = GetRoot()
        dossier = request.json.get('dossier')
        dossier_path = os.path.join(directory_path, dossier)
        if not dossier_path:
            return jsonify({'error': 'Missing dossier path'}), 400

        exists = os.path.exists(dossier_path)
        return jsonify({'exists': exists}), 200

    except Exception as e:
        logger.error(f"Error checking dossier existence: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    from multiprocessing import freeze_support
    from JO_analyse import *
    import webbrowser
    import threading
    import time
    
    def open_browser():
        """Ouvre le navigateur après un court délai"""
        time.sleep(1.5)  # Attendre que le serveur démarre
        webbrowser.open('http://127.0.0.1:5000')
    
    # Lancer l'ouverture du navigateur dans un thread séparé
    threading.Thread(target=open_browser).start()
    
    # Initialiser multiprocessing
    freeze_support()
    
    # Démarrer le serveur Flask (cette ligne est bloquante)
    app.run(debug=True)


